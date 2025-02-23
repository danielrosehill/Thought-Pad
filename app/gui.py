from PyQt5.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QPushButton,
    QComboBox, QTextEdit, QLineEdit, QLabel, QMessageBox, QProgressBar,
    QSplitter, QInputDialog, QStatusBar, QToolButton, QFrame, QDialog,
    QFormLayout, QMenuBar, QMenu, QAction, QFileDialog
)
from PyQt5.QtCore import Qt, QTimer, QThread, pyqtSignal, pyqtSlot
from PyQt5.QtGui import QIcon, QKeySequence, QTextCharFormat, QColor, QPalette
import pyqtgraph as pg
import numpy as np
import openai
from pathlib import Path
import threading
import queue
import time

from config import Config
from audio_manager import AudioManager

# Define color scheme
COLORS = {
    'primary': '#2196F3',    # Blue
    'secondary': '#4CAF50',  # Green
    'accent': '#FF9800',     # Orange
    'error': '#F44336',      # Red
    'background': '#F5F5F5', # Light Gray
    'text': '#212121',       # Dark Gray
    'disabled': '#9E9E9E'    # Medium Gray
}

class TranscribeWorker(QThread):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, api_key, audio_path):
        super().__init__()
        self.api_key = api_key
        self.audio_path = audio_path

    def run(self):
        try:
            if not Path(self.audio_path).exists():
                raise FileNotFoundError("Audio file not found")
                
            if not self.api_key:
                raise ValueError("OpenAI API key not set")

            client = openai.OpenAI(api_key=self.api_key)
            audio_file = Path(self.audio_path)
            
            with audio_file.open('rb') as f:
                transcript = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=f
                )
            self.finished.emit(transcript.text)
        except FileNotFoundError as e:
            self.error.emit(f"Audio file error: {str(e)}")
        except openai.AuthenticationError:
            self.error.emit("Invalid OpenAI API key")
        except openai.RateLimitError:
            self.error.emit("OpenAI API rate limit exceeded")
        except Exception as e:
            self.error.emit(f"Transcription error: {str(e)}")

class FormatWorker(QThread):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, api_key, text, temperature):
        super().__init__()
        self.api_key = api_key
        self.text = text
        self.temperature = temperature

    def run(self):
        try:
            client = openai.OpenAI(api_key=self.api_key)
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                temperature=self.temperature,
                messages=[
                    {"role": "system", "content": "Your task is to edit this text which was captured using speech to text. You should edit it lightly for clarity. You must remove any artifacts of spoken text, like 'um,' which you can assume the user did not intend to be captured in the finished document. If the user says things like 'take that out of the note' then you must use your reasoning to identify which parts of the text the user was referring to and remove those from the edited transcript. You must make sure that all thoughts and details are captured in the transcribed text. In addition to the formatted text, you must also suggest a title which captures the main essence of the note."},
                    {"role": "user", "content": self.text}
                ]
            )
            self.finished.emit(response.choices[0].message.content)
        except Exception as e:
            self.error.emit(str(e))

class SettingsDialog(QDialog):
    def __init__(self, config, audio_manager, parent=None):
        super().__init__(parent)
        self.config = config
        self.audio_manager = audio_manager
        self.setup_ui()

    def setup_ui(self):
        self.setWindowTitle("Settings")
        self.setMinimumWidth(400)
        layout = QFormLayout(self)

        # API Key
        self.api_key_input = QLineEdit(self.config.api_key)
        self.api_key_input.setEchoMode(QLineEdit.Password)
        show_key_btn = QToolButton()
        show_key_btn.setText("")
        show_key_btn.setToolTip("Show/Hide API Key")
        show_key_btn.setCheckable(True)
        show_key_btn.clicked.connect(self.toggle_api_key_visibility)
        
        api_key_layout = QHBoxLayout()
        api_key_layout.addWidget(self.api_key_input)
        api_key_layout.addWidget(show_key_btn)
        layout.addRow("OpenAI API Key:", api_key_layout)

        # Default Microphone
        self.mic_combo = QComboBox()
        self.populate_audio_devices()
        layout.addRow("Default Microphone:", self.mic_combo)

        # Default Export Format
        self.format_combo = QComboBox()
        self.format_combo.addItems(["Markdown", "PDF", "DocX", "Text"])
        current_format = self.config.get('default_export_format', 'Markdown')
        self.format_combo.setCurrentText(current_format)
        layout.addRow("Default Export Format:", self.format_combo)

        # Temperature Setting
        self.temp_combo = QComboBox()
        self.temp_combo.addItems(['0.0', '0.3', '0.5', '0.7', '1.0'])
        current_temp = str(self.config.get('gpt_temperature', 0.3))
        self.temp_combo.setCurrentText(current_temp)
        layout.addRow("GPT Temperature:", self.temp_combo)

        # Buttons
        button_box = QHBoxLayout()
        save_btn = QPushButton("Save")
        save_btn.clicked.connect(self.save_settings)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        button_box.addWidget(save_btn)
        button_box.addWidget(cancel_btn)
        layout.addRow("", button_box)

    def populate_audio_devices(self):
        devices = self.audio_manager.get_available_devices()
        self.mic_combo.clear()
        for device in devices:
            self.mic_combo.addItem(device['name'], device['index'])
        
        preferred_device = self.config.get_preferred_audio_device()
        if preferred_device is not None:
            index = self.mic_combo.findData(preferred_device)
            if index >= 0:
                self.mic_combo.setCurrentIndex(index)

    def toggle_api_key_visibility(self, checked):
        self.api_key_input.setEchoMode(
            QLineEdit.Normal if checked else QLineEdit.Password
        )

    def save_settings(self):
        # Save API Key
        api_key = self.api_key_input.text().strip()
        if api_key:
            self.config.api_key = api_key

        # Save preferred audio device
        device_index = self.mic_combo.currentData()
        self.config.set_preferred_audio_device(device_index)

        # Save default export format
        self.config.set('default_export_format', self.format_combo.currentText())

        # Save temperature
        self.config.set('gpt_temperature', float(self.temp_combo.currentText()))

        self.accept()

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.config = Config()
        self.audio_manager = AudioManager()
        self.waveform_data = np.zeros(1000)
        self.waveform_timer = QTimer()
        self.waveform_timer.timeout.connect(self.update_waveform)
        self.waveform_timer.start(50)  # Update every 50ms
        self.setup_menu()
        self.setup_ui()
        self.setup_auto_save()
        self.apply_stylesheet()
        self.check_api_key()

    def setup_menu(self):
        menubar = self.menuBar()
        
        # File Menu
        file_menu = menubar.addMenu('File')
        
        settings_action = QAction('Settings', self)
        settings_action.setShortcut('Ctrl+,')
        settings_action.triggered.connect(self.show_settings)
        file_menu.addAction(settings_action)
        
        file_menu.addSeparator()
        
        exit_action = QAction('Exit', self)
        exit_action.setShortcut('Ctrl+Q')
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

    def show_settings(self):
        dialog = SettingsDialog(self.config, self.audio_manager, self)
        if dialog.exec_() == QDialog.Accepted:
            # Refresh UI based on new settings
            self.populate_audio_devices()
            self.format_combo.setCurrentText(self.config.get('default_export_format', 'Markdown'))

    def setup_ui(self):
        self.setWindowTitle("Thought Pad")
        self.setMinimumSize(1200, 800)

        # Create main widget and layout
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QHBoxLayout(main_widget)

        # Create splitter for left and right panes
        splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(splitter)

        # Left pane (Audio controls)
        left_pane = QWidget()
        left_layout = QVBoxLayout(left_pane)
        
        # Audio device selection
        device_layout = QHBoxLayout()
        device_label = QLabel("Audio Input:")
        self.device_combo = QComboBox()
        self.populate_audio_devices()
        device_layout.addWidget(device_label)
        device_layout.addWidget(self.device_combo)
        left_layout.addLayout(device_layout)

        # Waveform display
        waveform_frame = QFrame()
        waveform_frame.setFrameStyle(QFrame.Panel | QFrame.Sunken)
        waveform_layout = QVBoxLayout(waveform_frame)
        
        # Create pyqtgraph plot widget
        pg.setConfigOptions(antialias=True)
        self.waveform_plot = pg.PlotWidget()
        self.waveform_plot.setBackground(COLORS['background'])
        self.waveform_plot.showGrid(x=True, y=True, alpha=0.3)
        self.waveform_plot.setYRange(-1, 1)
        self.waveform_plot.setTitle("Audio Waveform")
        self.waveform_curve = self.waveform_plot.plot(pen=pg.mkPen(color=COLORS['primary'], width=2))
        waveform_layout.addWidget(self.waveform_plot)
        left_layout.addWidget(waveform_frame)

        # Recording controls
        controls_layout = QHBoxLayout()
        self.record_button = QPushButton("Record")
        self.record_button.setToolTip("Start/Stop Recording (Ctrl+R)")
        self.record_button.setShortcut(QKeySequence("Ctrl+R"))
        
        self.pause_button = QPushButton("Pause")
        self.pause_button.setToolTip("Pause/Resume Recording (Ctrl+P)")
        self.pause_button.setShortcut(QKeySequence("Ctrl+P"))
        self.pause_button.setEnabled(False)
        
        self.stop_button = QPushButton("Stop")
        self.stop_button.setToolTip("Stop Recording and Transcribe (Ctrl+S)")
        self.stop_button.setShortcut(QKeySequence("Ctrl+S"))
        self.stop_button.setEnabled(False)
        
        self.clear_button = QPushButton("Clear")
        self.clear_button.setToolTip("Clear All Text (Ctrl+L)")
        self.clear_button.setShortcut(QKeySequence("Ctrl+L"))
        
        controls_layout.addWidget(self.record_button)
        controls_layout.addWidget(self.pause_button)
        controls_layout.addWidget(self.stop_button)
        controls_layout.addWidget(self.clear_button)
        
        left_layout.addLayout(controls_layout)

        # Recording status indicator
        self.status_label = QLabel("Ready")
        self.status_label.setAlignment(Qt.AlignCenter)
        self.status_label.setStyleSheet(f"color: {COLORS['secondary']}")
        left_layout.addWidget(self.status_label)

        # Right pane (Text display and processing)
        right_pane = QWidget()
        right_layout = QVBoxLayout(right_pane)

        # Title field
        title_layout = QHBoxLayout()
        title_label = QLabel("Title:")
        self.title_edit = QLineEdit()
        self.title_edit.setPlaceholderText("Enter note title here")
        title_layout.addWidget(title_label)
        title_layout.addWidget(self.title_edit)
        right_layout.addLayout(title_layout)

        # Raw transcribed text
        raw_layout = QHBoxLayout()
        raw_label = QLabel("Transcribed Text:")
        self.edit_mode_btn = QToolButton()
        self.edit_mode_btn.setToolTip("Toggle Edit Mode (Ctrl+E)")
        self.edit_mode_btn.setText("")
        self.edit_mode_btn.setCheckable(True)
        self.edit_mode_btn.clicked.connect(self.toggle_edit_mode)
        raw_layout.addWidget(raw_label)
        raw_layout.addWidget(self.edit_mode_btn)
        raw_layout.addStretch()
        self.word_count_label = QLabel("Words: 0")
        raw_layout.addWidget(self.word_count_label)
        right_layout.addLayout(raw_layout)
        
        self.raw_text = QTextEdit()
        self.raw_text.setReadOnly(True)
        self.raw_text.textChanged.connect(self.update_word_count)
        right_layout.addWidget(self.raw_text)

        # Formatted text
        formatted_label = QLabel("Formatted Text:")
        self.formatted_text = QTextEdit()
        right_layout.addWidget(formatted_label)
        right_layout.addWidget(self.formatted_text)

        # Bottom controls
        bottom_layout = QHBoxLayout()
        self.format_button = QPushButton("Format Text")
        self.format_button.setToolTip("Format Text (Ctrl+F)")
        self.format_button.setShortcut(QKeySequence("Ctrl+F"))
        
        self.download_button = QPushButton("Download All")
        self.download_button.setToolTip("Download Note (Ctrl+D)")
        self.download_button.setShortcut(QKeySequence("Ctrl+D"))
        
        self.format_combo = QComboBox()
        self.format_combo.addItems(["Markdown", "PDF", "DocX", "Text"])
        default_format = self.config.get('default_export_format', 'Markdown')
        self.format_combo.setCurrentText(default_format)
        
        bottom_layout.addWidget(self.format_button)
        bottom_layout.addWidget(self.format_combo)
        bottom_layout.addWidget(self.download_button)
        right_layout.addLayout(bottom_layout)

        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.hide()
        right_layout.addWidget(self.progress_bar)

        # Add panes to splitter
        splitter.addWidget(left_pane)
        splitter.addWidget(right_pane)
        splitter.setSizes([400, 800])  # 1:2 split

        # Add status bar
        self.statusBar = QStatusBar()
        self.setStatusBar(self.statusBar)
        self.statusBar.showMessage("Ready")

        # Connect signals
        self.connect_signals()

    def connect_signals(self):
        self.record_button.clicked.connect(self.toggle_recording)
        self.pause_button.clicked.connect(self.toggle_pause)
        self.stop_button.clicked.connect(self.stop_recording)
        self.clear_button.clicked.connect(self.clear_all)
        self.format_button.clicked.connect(self.format_text)
        self.download_button.clicked.connect(self.download_text)
        self.device_combo.currentIndexChanged.connect(self.change_audio_device)

    def populate_audio_devices(self):
        devices = self.audio_manager.get_available_devices()
        self.device_combo.clear()
        for device in devices:
            self.device_combo.addItem(device['name'], device['index'])
        
        # Set preferred device if configured
        preferred_device = self.config.get_preferred_audio_device()
        if preferred_device is not None:
            index = self.device_combo.findData(preferred_device)
            if index >= 0:
                self.device_combo.setCurrentIndex(index)

    def change_audio_device(self, index):
        device_index = self.device_combo.currentData()
        if device_index is not None:
            self.audio_manager.set_device(device_index)
            self.config.set_preferred_audio_device(device_index)

    def toggle_recording(self):
        if not self.audio_manager.recording:
            self.audio_manager.start_recording()
            self.record_button.setText("Stop Recording")
            self.pause_button.setEnabled(True)
            self.stop_button.setEnabled(True)
            self.update_status("Recording...", COLORS['error'])
            self.waveform_data = np.zeros(1000)
        else:
            self.stop_recording()

    def toggle_pause(self):
        if self.audio_manager.paused:
            self.audio_manager.resume_recording()
            self.update_status("Recording...", COLORS['error'])
        else:
            self.audio_manager.pause_recording()
            self.update_status("Paused", COLORS['accent'])

    def stop_recording(self):
        self.audio_manager.stop_recording()
        self.record_button.setText("Record")
        self.pause_button.setEnabled(False)
        self.stop_button.setEnabled(False)
        self.update_status("Transcribing...", COLORS['primary'])
        self.transcribe_audio()

    def on_transcription_complete(self, text):
        self.raw_text.setText(text)
        self.update_status("Ready", COLORS['secondary'])
        self.progress_bar.hide()

    def on_transcription_error(self, error):
        self.show_error("Transcription Error", str(error))
        self.update_status("Ready", COLORS['secondary'])
        self.progress_bar.hide()

    def transcribe_audio(self):
        if not self.config.api_key:
            self.check_api_key()
            if not self.config.api_key:
                return

        audio_file = self.audio_manager.get_temp_file_path()
        if not audio_file:
            self.show_error("Error", "No audio recorded")
            self.update_status("Ready", COLORS['secondary'])
            return

        self.progress_bar.show()
        self.worker = TranscribeWorker(self.config.api_key, audio_file)
        self.worker.finished.connect(self.on_transcription_complete)
        self.worker.error.connect(self.on_transcription_error)
        self.worker.start()

    def format_text(self):
        if not self.raw_text.toPlainText():
            QMessageBox.warning(self, "Error", "No text to format!")
            return
        self.progress_bar.setRange(0, 0)
        self.progress_bar.show()

        self.format_worker = FormatWorker(
            self.config.api_key,
            self.raw_text.toPlainText(),
            self.config.get('gpt_temperature', 0.3)
        )
        self.format_worker.finished.connect(self._on_format_finished)
        self.format_worker.error.connect(self._on_format_error)
        self.format_worker.start()

    @pyqtSlot(str)
    def _on_format_finished(self, formatted_text):
        lines = formatted_text.split('\n')
        title = lines[0].replace('Title: ', '').strip()
        content = '\n'.join(lines[1:]).strip()
        self.title_edit.setText(title)
        self.formatted_text.setText(content)
        self.progress_bar.hide()

    @pyqtSlot(str)
    def _on_format_error(self, error):
        self.show_error("Text formatting failed", error)
        self.progress_bar.hide()

    def download_text(self):
        if not self.formatted_text.toPlainText():
            QMessageBox.warning(self, "Error", "Please format the text before downloading!")
            return

        title = self.title_edit.text() or "untitled"
        format_type = self.format_combo.currentText()
        
        try:
            if format_type == "Text":
                file_path, _ = QFileDialog.getSaveFileName(
                    self, "Save Text File", "", "Text Files (*.txt)"
                )
                if file_path:
                    with open(file_path, 'w') as f:
                        f.write(self.formatted_text.toPlainText())
            elif format_type == "Markdown":
                self.save_markdown(title)
            elif format_type == "PDF":
                self.save_pdf(title)
            else:  # DocX
                self.save_docx(title)
                
            self.statusBar.showMessage(f"File saved successfully as {format_type}")
        except Exception as e:
            self.show_error("Save Error", f"Error saving file: {str(e)}")

    def save_markdown(self, title):
        content = (f"# {title}\n\n"
                  f"{self.formatted_text.toPlainText()}\n\n"
                  f"---FORMATTED---\n\n"
                  f"{self.raw_text.toPlainText()}\n\n"
                  f"---RAW---")
        
        path = Path(f"{title}.md")
        try:
            path.write_text(content)
            QMessageBox.information(self, "Success", f"File saved as {path}")
        except Exception as e:
            self.show_error("Error saving file", str(e))

    def save_pdf(self, title):
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Add title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, title, ln=True)
        pdf.ln(10)
        
        # Add formatted text
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, self.formatted_text.toPlainText())
        pdf.ln(10)
        
        # Add separator and raw text
        pdf.cell(0, 10, "---FORMATTED---", ln=True)
        pdf.ln(10)
        pdf.multi_cell(0, 10, self.raw_text.toPlainText())
        pdf.cell(0, 10, "---RAW---", ln=True)
        
        try:
            pdf.output(f"{title}.pdf")
            QMessageBox.information(self, "Success", f"File saved as {title}.pdf")
        except Exception as e:
            self.show_error("Error saving PDF", str(e))

    def save_docx(self, title):
        from docx import Document
        doc = Document()
        
        # Add title
        doc.add_heading(title, 0)
        
        # Add formatted text
        doc.add_paragraph(self.formatted_text.toPlainText())
        
        # Add separator and raw text
        doc.add_paragraph("---FORMATTED---")
        doc.add_paragraph(self.raw_text.toPlainText())
        doc.add_paragraph("---RAW---")
        
        try:
            doc.save(f"{title}.docx")
            QMessageBox.information(self, "Success", f"File saved as {title}.docx")
        except Exception as e:
            self.show_error("Error saving DOCX", str(e))

    def clear_all(self):
        self.audio_manager.clear_recording()
        self.raw_text.clear()
        self.formatted_text.clear()
        self.title_edit.clear()

    def setup_auto_save(self):
        self.auto_save_timer = QTimer()
        self.auto_save_timer.timeout.connect(self.auto_save)
        interval = self.config.get('auto_save_interval', 300) * 1000  # Convert to milliseconds
        self.auto_save_timer.start(interval)

    def auto_save(self):
        # TODO: Implement auto-save functionality
        pass

    def check_api_key(self):
        """Check if OpenAI API key is set and prompt user if not."""
        api_key = self.config.api_key
        if not api_key:
            self.show_settings()
            api_key = self.config.api_key
            if not api_key:
                self.show_error(
                    "API Key Required",
                    "Please set your OpenAI API key in the settings (Ctrl+,)"
                )
        return bool(api_key)

    def show_error(self, title, message):
        QMessageBox.critical(self, title, message)

    def toggle_edit_mode(self):
        is_editable = self.edit_mode_btn.isChecked()
        self.raw_text.setReadOnly(not is_editable)
        self.edit_mode_btn.setText("" if is_editable else "")
        self.statusBar.showMessage("Edit mode " + ("enabled" if is_editable else "disabled"))

    def update_word_count(self):
        text = self.raw_text.toPlainText()
        word_count = len(text.split()) if text else 0
        self.word_count_label.setText(f"Words: {word_count}")

    def update_status(self, message, color="black"):
        self.status_label.setText(message)
        self.status_label.setStyleSheet(f"color: {color}")
        self.statusBar.showMessage(message)

    def update_waveform(self):
        if hasattr(self.audio_manager, 'recorded_frames') and self.audio_manager.recording:
            with self.audio_manager.frames_lock:
                if len(self.audio_manager.recorded_frames) > 0:
                    # Get the latest frame
                    latest_frame = self.audio_manager.recorded_frames[-1]
                    # Roll the existing data and add new points
                    self.waveform_data = np.roll(self.waveform_data, -len(latest_frame))
                    self.waveform_data[-len(latest_frame):] = latest_frame.flatten()
                    self.waveform_curve.setData(self.waveform_data)

    def apply_stylesheet(self):
        self.setStyleSheet(f"""
            QMainWindow {{
                background-color: {COLORS['background']};
            }}
            QPushButton {{
                background-color: {COLORS['primary']};
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
            }}
            QPushButton:hover {{
                background-color: {COLORS['secondary']};
            }}
            QPushButton:disabled {{
                background-color: {COLORS['disabled']};
            }}
            QLabel {{
                color: {COLORS['text']};
            }}
            QTextEdit {{
                border: 1px solid {COLORS['disabled']};
                border-radius: 4px;
                padding: 4px;
            }}
            QComboBox {{
                border: 1px solid {COLORS['disabled']};
                border-radius: 4px;
                padding: 4px;
            }}
        """)